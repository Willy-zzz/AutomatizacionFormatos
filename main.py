import pandas as pd
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import re
import os


# Función mejorada para reemplazar texto en párrafos y tablas
def reemplazar_campos(doc, context):
    def reemplazar_en_parrafos(parrafos):
        for p in parrafos:
            texto_completo = ''.join(run.text for run in p.runs)
            for key, value in context.items():
                if f'{{{{{key}}}}}' in texto_completo:
                    nuevo_texto = texto_completo.replace(f'{{{{{key}}}}}', str(value))
                    # Borrar runs anteriores
                    for run in p.runs:
                        run.text = ''
                    # Insertar nuevo texto como un solo run
                    if p.runs:
                        p.runs[0].text = nuevo_texto

    # Reemplazar en párrafos fuera de tablas
    reemplazar_en_parrafos(doc.paragraphs)

    # Reemplazar dentro de cada celda de tabla
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_en_parrafos(cell.paragraphs)

# Cargar Excel y extraer IDs
def cargar_excel():
    archivo = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if archivo:
        global df
        df = pd.read_excel(archivo, header=[0,1])  # lee dos filas como encabezado
        # Aplanar columnas, quitar 'Unnamed...' del segundo nivel
        df.columns = [' '.join([str(i) for i in col if 'Unnamed' not in str(i)]).strip() for col in df.columns.values]
        print("CAMPOS DISPONIBLES EN EXCEL:")
        print(df.columns.tolist())  # Verifica nombres limpios
        combo['values'] = df['ID / Posición'].astype(str).unique().tolist()
        messagebox.showinfo("Éxito", "Archivo cargado. Ahora selecciona un ID.")

# Generar documento Word con datos
def generar_word():
    try:
        id_valor = combo.get()
        if not id_valor:
            raise Exception("Selecciona un ID / Posición.")
        
        datos = df[df['ID / Posición'].astype(str) == id_valor].iloc[0]

        print("FILTRANDO DATOS PARA ID:", id_valor)
        print("DATOS ENCONTRADOS:")
        print(datos)

        # Diccionario para reemplazo
        context = {
            'NumeroPlaza': datos['Número de plaza:'],
            'IdPosicion': datos['ID / Posición'],
            'Nivel': datos['Nivel:'],
            'Grupo': datos['Grupo:'],
            'EPSoDireccion': datos['EPS o Dirección Corporativa:'],
            'Subdireccion': datos['Subdirección'],
            'CentroTrabajo': datos['Centro de Trabajo:'],
            'DenominacionPuesto': datos['Denominación de puesto:'],
            'MotivoDisponibilidada': datos['Motivo o justificación de disponibilidad:'],
            'UltimoOcupante': datos['Último ocupante:'],
            'TipoContrato': datos['Tipo de contrato:'],
            'DisponibilidadApartir': str(datos['Disponible a partir de:']),
            'Categoria': datos['Categoría:'],
            'Clasificacion': datos['Clasificación:'],
            'RequerimientoProfesional': datos['Requerimiento profesional:'],
            'NumeroPuesto': datos['Número de Puesto:'],
            'Jornada': datos['Jornada:'],

            # Candidato (basado en CANDIDATO 1)
            'Nombre': datos['CANDIDATO 1 Nombre:'],
            'IdFicha': datos['CANDIDATO 1 ID BT/Ficha:'],
            'RegimenContractual': datos['CANDIDATO 1 Régimen Contractual:'],
            'NivelC': datos['CANDIDATO 1 Nivel:'],
            'GrupoPActual': datos['CANDIDATO 1 Grupo plaza actual:'],
            'EpsCorporativa': datos['CANDIDATO 1 EPS o Dirección Corporativa:'],
            'SubdireccionC': datos['CANDIDATO 1 Subdirección:'],
            'CentroTrabajoC': datos['CANDIDATO 1 Centro de trabajo:'],
            'Departamento': datos['CANDIDATO 1 Departamento'],
            'PuestoActual': datos['CANDIDATO 1 Puesto Actual:'],
            'Formacion': datos['CANDIDATO 1 Formación:'],
            'Cedula': datos['CANDIDATO 1 No. Ced.'],

            # Extra
            'Observaciones': datos['Observaciones:'],
            'ResponsableValidacion': datos['Responsable de la validación:'],
            'CargoResponsable': datos['Cargo Responsable de la validación:'],
            'CandidatoSeleccionado': datos['Candidato Seleccionado:'],
            'JustificacionSeleccion': datos['Justificación de la selección:'],
            'ResponsableSeleccion': datos['Responsable de la Selección:'],
            'CargoSeleccion': datos['Cargo Responsable de la Selección:']
        }

        # Cargar plantilla Word
        plantilla = Document('Plantillas\Plantilla_FSV.docx')
        print("VALORES QUE SE VAN A RELLENAR EN WORD:")
        for k, v in context.items():
            print(f"{k}: {v}")

        reemplazar_campos(plantilla, context)

        # Verificar si quedaron etiquetas sin reemplazar
        print("⚠️ VERIFICANDO CAMPOS PENDIENTES:")
        for p in plantilla.paragraphs:
            if '{{' in p.text:
                print("⚠️ En párrafo:", p.text)
        for table in plantilla.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '{{' in p.text:
                            print("⚠️ En tabla:", p.text)

         # Pedir al usuario dónde guardar
        nombre_candidato = datos['Candidato Seleccionado:']
        nombre_archivo = re.sub(r'[\\/*?:"<>|]', "", nombre_candidato)
        ruta_guardado = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")],
            initialfile=f"FVS {nombre_archivo}.docx",
            title="Guardar documento como..."
        )

        if ruta_guardado:
            plantilla.save(ruta_guardado)
            messagebox.showinfo("Éxito", f"Documento guardado en:\n{ruta_guardado}")
        else:
            messagebox.showwarning("Cancelado", "No se guardó el documento.")
    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un error:\n{e}")

# GUI principal
root = tk.Tk()
root.title("Generador FVS")
root.geometry("400x250")

tk.Label(root, text="1. Cargar archivo Excel").pack(pady=5)
tk.Button(root, text="Seleccionar Excel", command=cargar_excel).pack()

tk.Label(root, text="2. Seleccionar ID / Posición").pack(pady=5)
combo = ttk.Combobox(root, state="readonly")
combo.pack()

tk.Label(root, text="3. Generar documento Word").pack(pady=5)
tk.Button(root, text="Generar FVS", command=generar_word).pack(pady=10)

root.mainloop()
